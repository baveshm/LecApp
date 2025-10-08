"""
Utility functions for LecApp.

This module contains self-contained helper functions used throughout the application
for JSON processing, HTML/Markdown conversion, and document formatting.
"""

import json
import re
import ast
import markdown
import bleach
from datetime import datetime, timedelta
import uuid


# --- JSON Processing Functions ---

def auto_close_json(json_string):
    """
    Attempts to close an incomplete JSON string by appending necessary brackets and braces.
    This is a simplified parser and may not handle all edge cases, but is
    designed to fix unterminated strings from API responses.
    """
    if not isinstance(json_string, str):
        return json_string

    stack = []
    in_string = False
    escape_next = False

    for char in json_string:
        if escape_next:
            escape_next = False
            continue

        if char == '\\':
            escape_next = True
            continue

        if char == '"':
            # We don't handle escaped quotes inside strings perfectly,
            # but this is a simple heuristic.
            if not escape_next:
                in_string = not in_string

        if not in_string:
            if char == '{':
                stack.append('}')
            elif char == '[':
                stack.append(']')
            elif char == '}':
                if stack and stack[-1] == '}':
                    stack.pop()
            elif char == ']':
                if stack and stack[-1] == ']':
                    stack.pop()

    # If we are inside a string at the end, close it.
    if in_string:
        json_string += '"'
    
    # Close any remaining open structures
    while stack:
        json_string += stack.pop()

    return json_string


def preprocess_json_escapes(json_string):
    """
    Preprocess JSON string to fix common escape issues from LLM responses.
    Uses a more sophisticated approach to handle nested quotes properly.
    """
    if not json_string:
        return json_string
    
    result = []
    i = 0
    in_string = False
    escape_next = False
    expecting_value = False  # Track if we're expecting a value (after :)
    
    while i < len(json_string):
        char = json_string[i]
        
        if escape_next:
            # This character is escaped, add it as-is
            result.append(char)
            escape_next = False
        elif char == '\\':
            # This is an escape character
            result.append(char)
            escape_next = True
        elif char == ':' and not in_string:
            # We found a colon, next string will be a value
            result.append(char)
            expecting_value = True
        elif char == ',' and not in_string:
            # We found a comma, reset expecting_value
            result.append(char)
            expecting_value = False
        elif char == '"':
            if not in_string:
                # Starting a string
                in_string = True
                result.append(char)
            else:
                # Determine if this is a terminating quote by peeking ahead
                j = i + 1
                while j < len(json_string) and json_string[j].isspace():
                    j += 1
                # Keys end with ':' outside arrays/objects; values end with , } ] or end of string
                if expecting_value:
                    terminators = ',}]'
                else:
                    terminators = ':'
                if j >= len(json_string) or (j < len(json_string) and json_string[j] in terminators):
                    # End of string
                    in_string = False
                    result.append(char)
                    if not expecting_value:
                        expecting_value = True
                else:
                    # Inner quote -> escape
                    result.append('\\"')
        elif in_string and char in ('\n', '\r'):
            # Escape raw newlines inside JSON strings
            if char == '\n':
                result.append('\\n')
            else:
                result.append('\\r')
        else:
            result.append(char)
        
        i += 1
    
    return ''.join(result)


def extract_json_object(text):
    """Extract the first balanced JSON object or array from arbitrary text.

    Uses a lightweight state machine to find the earliest '{' or '[' and
    then walks forward tracking nesting depth while respecting quoted
    strings and escape characters. Returns the substring containing the
    balanced JSON structure, or the original text if none found.
    """
    if not text or not isinstance(text, str):
        return text

    start_indexes = []
    for idx, ch in enumerate(text):
        if ch in '{[':
            start_indexes.append(idx)
            break  # Only need the first structure

    if not start_indexes:
        return text

    start = start_indexes[0]
    opening = text[start]
    closing = '}' if opening == '{' else ']'
    depth = 0
    in_string = False
    escape = False

    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if c == '\\':
            escape = True
            continue
        if c == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if c == opening:
            depth += 1
        elif c == closing:
            depth -= 1
            if depth == 0:
                return text[start:i+1]
    # Fallback: return original text if not balanced
    return text


def safe_json_loads(json_string, fallback_value=None, logger=None):
    """
    Safely parse JSON with preprocessing to handle common LLM JSON formatting issues.
    
    Args:
        json_string (str): The JSON string to parse
        fallback_value: Value to return if parsing fails (default: None)
        logger: Optional logger instance for logging warnings/errors
    
    Returns:
        Parsed JSON object or fallback_value if parsing fails
    """
    if not json_string or not isinstance(json_string, str):
        if logger:
            logger.warning(f"Invalid JSON input: {type(json_string)} - {json_string}")
        return fallback_value
    
    # Step 1: Clean the input string
    cleaned_json = json_string.strip()
    
    # Step 2: Extract JSON from markdown code blocks if present
    json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', cleaned_json, re.DOTALL)
    if json_match:
        cleaned_json = json_match.group(1).strip()
    
    # Step 3: Try multiple parsing strategies
    parsing_strategies = [
        # Strategy 1: Direct parsing (for well-formed JSON)
        lambda x: json.loads(x),

        # Strategy 1b: Sanitize raw newlines and inner quotes inside value strings (relaxed JSON)
        lambda x: json.loads(sanitize_relaxed_json(x)),

        # Strategy 2: Extract balanced JSON substring then parse (helps when wrapped in text)
        lambda x: json.loads(extract_json_object(x)),

        # Strategy 3: Fix common escape issues
        lambda x: json.loads(preprocess_json_escapes(x)),

        # Strategy 4: Use ast.literal_eval as fallback for simple cases
        lambda x: ast.literal_eval(x) if x.startswith(('{', '[')) else None,

        # Strategy 5: Auto-close incomplete JSON and parse
        lambda x: json.loads(auto_close_json(x)),

        # Strategy 6: Aggressive inner quote escaping within value strings
        lambda x: json.loads(_escape_inner_quotes(x)),

        # Strategy 7: Robust multiline value quote escaper
        lambda x: json.loads(robust_escape_value_quotes(x)),

        # Strategy 8: Regex-based pair escaping inside values
        lambda x: json.loads(regex_escape_inner_pairs(x)),
    ]
    
    for i, strategy in enumerate(parsing_strategies):
        try:
            result = strategy(cleaned_json)
            if result is not None:
                if i > 0 and logger:  # Log if we had to use a fallback strategy
                    logger.info(f"JSON parsed successfully using strategy {i+1}")
                return result
        except (json.JSONDecodeError, ValueError, SyntaxError) as e:
            if i == 0 and logger:  # Only log the first failure to avoid spam
                logger.debug(f"JSON parsing strategy {i+1} failed: {e}")
            continue
    
    # All strategies failed
    if logger:
        logger.error(f"All JSON parsing strategies failed for: {cleaned_json[:200]}...")
    return fallback_value


def _escape_inner_quotes(s):
    """Attempt to escape problematic inner quotes inside JSON string values.

    This heuristic scans for sequences like "... unescaped "word" ..." inside
    a value string and escapes the inner quotes. It avoids touching already
    escaped quotes and keys.
    """
    if not s or '"' not in s:
        return s
    result = []
    in_string = False
    escape = False
    for i, ch in enumerate(s):
        if escape:
            result.append(ch)
            escape = False
            continue
        if ch == '\\':
            result.append(ch)
            escape = True
            continue
        if ch == '"':
            # Toggle string context
            in_string = not in_string
            result.append(ch)
            continue
        if ch == '"' and in_string:
            # Inner quote inside a string (unlikely since handled above)
            result.append('\\"')
            continue
        if in_string and ch == '"':
            result.append('\\"')
        else:
            result.append(ch)
    return ''.join(result)


def robust_escape_value_quotes(s):
    """Escapes unescaped quotes inside JSON string values, including multiline values.

    Walks through the JSON text and when inside a string, any quote that would
    prematurely terminate the value (i.e., not followed by optional whitespace
    then a structural delimiter , } ] ) is escaped.
    """
    if not s or '"' not in s:
        return s
    chars = []
    in_string = False
    escape = False
    for i, ch in enumerate(s):
        if escape:
            chars.append(ch)
            escape = False
            continue
        if ch == '\\':
            chars.append(ch)
            escape = True
            continue
        if ch == '"':
            if not in_string:
                in_string = True
                chars.append(ch)
            else:
                # Peek ahead
                j = i + 1
                while j < len(s) and s[j].isspace():
                    j += 1
                if j >= len(s) or s[j] in ',}]':
                    # Legitimate terminator
                    in_string = False
                    chars.append(ch)
                else:
                    # Escape this inner quote
                    chars.append('\\"')
        else:
            chars.append(ch)
    return ''.join(chars)


def regex_escape_inner_pairs(s):
    """Use regex to escape unescaped quote pairs inside JSON string values.

    For each JSON string value, finds occurrences of unescaped "word" (without
    structural termination) and escapes the quotes. This acts as a last resort
    and only runs if previous strategies failed.
    """
    if not s or '"' not in s:
        return s
    # Pattern to find JSON string values: "..." including multiline (DOTALL)
    string_pattern = re.compile(r'"((?:[^"\\]|\\.)*)"', re.DOTALL)

    def fix_content(match):
        content = match.group(1)
        # Escape inner unescaped quote pairs not followed by structural delimiters
        # Replace "word" where word has no quotes inside
        inner_pattern = re.compile(r'(?<!\\)"([^"\n]{1,80}?)"')
        # Apply repeatedly until no change (avoid infinite loops with proper condition)
        prev = None
        while prev != content:
            prev = content
            content = inner_pattern.sub(r'\\"\1\\"', content)
        return '"' + content + '"'
    try:
        return string_pattern.sub(fix_content, s)
    except re.error:
        return s


def sanitize_relaxed_json(s):
    """Convert a relaxed JSON-like string containing raw newlines and unescaped
    inner quotes inside value strings into valid JSON.

    - Replaces raw newline characters within strings with \n
    - Escapes embedded quotes that are not string terminators
    """
    if not s or '"' not in s:
        return s
    chars = []
    in_string = False
    escape = False
    for i, ch in enumerate(s):
        if escape:
            chars.append(ch)
            escape = False
            continue
        if ch == '\\':
            chars.append(ch)
            escape = True
            continue
        if ch == '"':
            if not in_string:
                in_string = True
                chars.append(ch)
            else:
                # Peek ahead to decide termination
                j = i + 1
                while j < len(s) and s[j].isspace() and s[j] not in '\n\r':
                    j += 1
                if j >= len(s) or s[j] in ',}]':
                    in_string = False
                    chars.append(ch)
                else:
                    # Inner quote
                    chars.append('\\"')
        elif ch in ('\n', '\r') and in_string:
            # Encode newline inside string
            if ch == '\n':
                chars.append('\\n')
            else:
                chars.append('\\r')
        else:
            chars.append(ch)
    return ''.join(chars)


# --- HTML/Markdown Processing Functions ---

def sanitize_html(text):
    """
    Sanitize HTML content to prevent XSS attacks and code execution.
    This function removes dangerous content while preserving safe formatting.
    """
    if not text:
        return ""
    
    # First, remove any template syntax that could be dangerous
    # Remove Jinja2/Flask template syntax
    text = re.sub(r'\{\{.*?\}\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\{%.*?%\}', '', text, flags=re.DOTALL)
    
    # Remove other template-like syntax
    text = re.sub(r'<%.*?%>', '', text, flags=re.DOTALL)
    text = re.sub(r'<\?.*?\?>', '', text, flags=re.DOTALL)
    
    # Define allowed tags and attributes for safe HTML
    allowed_tags = [
        'p', 'br', 'strong', 'b', 'em', 'i', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'li', 'blockquote', 'code', 'pre', 'a', 'img', 'table', 'thead', 
        'tbody', 'tr', 'th', 'td', 'dl', 'dt', 'dd', 'div', 'span', 'hr', 'sup', 'sub'
    ]
    
    allowed_attributes = {
        'a': ['href', 'title'],
        'img': ['src', 'alt', 'title', 'width', 'height'],
        'code': ['class'],  # For syntax highlighting
        'pre': ['class'],   # For syntax highlighting
        'div': ['class'],   # For code blocks
        'span': ['class'],  # For syntax highlighting
        'th': ['align'],
        'td': ['align'],
        'table': ['class']
    }
    
    # Sanitize the HTML to remove dangerous content
    sanitized_html = bleach.clean(
        text,
        tags=allowed_tags,
        attributes=allowed_attributes,
        protocols=['http', 'https', 'mailto'],
        strip=True  # Strip disallowed tags instead of escaping them
    )
    
    return sanitized_html


def md_to_html(text):
    """
    Convert markdown text to HTML with proper formatting and sanitization.
    
    Args:
        text (str): Markdown text to convert
        
    Returns:
        str: Sanitized HTML output
    """
    if not text:
        return ""
    
    # Fix list spacing
    def fix_list_spacing(text):
        lines = text.split('\n')
        result = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Check if this line is a list item (starts with -, *, +, or number.)
            is_list_item = (
                stripped.startswith(('- ', '* ', '+ ')) or
                (stripped and stripped[0].isdigit() and '. ' in stripped[:10])
            )
            
            # If we're starting a new list or continuing a list, ensure proper spacing
            if is_list_item:
                if not in_list and result and result[-1].strip():
                    # Starting a new list - add blank line before
                    result.append('')
                in_list = True
            elif in_list and stripped and not is_list_item:
                # Ending a list - add blank line after the list
                if result and result[-1].strip():
                    result.append('')
                in_list = False
            
            result.append(line)
        
        return '\n'.join(result)
    
    # Fix list spacing
    processed_text = fix_list_spacing(text)
    
    # Convert markdown to HTML with extensions for tables, code highlighting, etc.
    html = markdown.markdown(processed_text, extensions=[
        'fenced_code',      # Fenced code blocks
        'tables',           # Table support
        'attr_list',        # Attribute lists
        'def_list',         # Definition lists
        'footnotes',        # Footnotes
        'abbr',             # Abbreviations
        'codehilite',       # Syntax highlighting for code blocks
        'smarty'            # Smart quotes, dashes, etc.
    ])
    
    # Apply sanitization to the generated HTML
    return sanitize_html(html)


def format_transcription_for_llm(transcription_text):
    """
    Formats transcription for LLM. If it's our simplified JSON, convert it to plain text.
    Otherwise, return as is.
    """
    try:
        transcription_data = json.loads(transcription_text)
        if isinstance(transcription_data, list):
            # It's our simplified JSON format
            formatted_lines = []
            for segment in transcription_data:
                speaker = segment.get('speaker', 'Unknown Speaker')
                sentence = segment.get('sentence', '')
                formatted_lines.append(f"[{speaker}]: {sentence}")
            return "\n".join(formatted_lines)
    except (json.JSONDecodeError, TypeError):
        # Not a JSON, or not the format we expect, so return as is.
        pass
    return transcription_text


def clean_llm_response(text):
    """
    Clean LLM responses by removing thinking tags and excessive whitespace.
    This handles responses from reasoning models that include <think> tags.
    """
    if not text:
        return ""
    
    # Remove thinking tags and their content
    # Handle both <think> and <thinking> tags with various closing formats
    cleaned = re.sub(r'<think(?:ing)?>.*?</think(?:ing)?>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Also handle unclosed thinking tags (in case the model doesn't close them)
    cleaned = re.sub(r'<think(?:ing)?>.*$', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove any remaining XML-like tags that might be related to thinking
    # but preserve markdown formatting
    cleaned = re.sub(r'<(?!/?(?:code|pre|blockquote|p|br|hr|ul|ol|li|h[1-6]|em|strong|b|i|a|img)(?:\s|>|/))[^>]+>', '', cleaned)
    
    # Clean up excessive whitespace while preserving intentional formatting
    # Remove leading/trailing whitespace from each line
    lines = cleaned.split('\n')
    cleaned_lines = []
    for line in lines:
        # Preserve lines that are part of code blocks or lists
        if line.strip() or (len(cleaned_lines) > 0 and cleaned_lines[-1].strip().startswith(('```', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
            cleaned_lines.append(line.rstrip())
    
    # Join lines and remove multiple consecutive blank lines
    cleaned = '\n'.join(cleaned_lines)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    
    # Final strip to remove leading/trailing whitespace
    return cleaned.strip()


def extract_thinking_content(text):
    """
    Extract thinking content from LLM responses.
    Returns a tuple of (thinking_content, main_content).
    """
    if not text:
        return ("", "")
    
    # Find all thinking tags and their content
    thinking_pattern = r'<think(?:ing)?>.*?</think(?:ing)?>'
    thinking_matches = re.findall(thinking_pattern, text, flags=re.DOTALL | re.IGNORECASE)
    
    # Extract the content from within the tags
    thinking_content = ""
    for match in thinking_matches:
        # Remove the opening and closing tags
        content = re.sub(r'^<think(?:ing)?>', '', match, flags=re.IGNORECASE)
        content = re.sub(r'</think(?:ing)?>$', '', content, flags=re.IGNORECASE)
        if thinking_content:
            thinking_content += "\n\n"
        thinking_content += content.strip()
    
    # Get the main content by removing thinking tags
    main_content = clean_llm_response(text)
    
    return (thinking_content, main_content)


def process_streaming_with_thinking(stream):
    """
    Generator that processes a streaming response and separates thinking content.
    Yields SSE-formatted data with 'delta' for regular content and 'thinking' for thinking content.
    """
    content_buffer = ""
    in_thinking = False
    thinking_buffer = ""
    
    for chunk in stream:
        content = chunk.choices[0].delta.content
        if content:
            content_buffer += content
            
            # Process the buffer to detect and handle thinking tags
            while True:
                if not in_thinking:
                    # Look for opening thinking tag
                    think_start = re.search(r'<think(?:ing)?>', content_buffer, re.IGNORECASE)
                    if think_start:
                        # Send any content before the thinking tag
                        before_thinking = content_buffer[:think_start.start()]
                        if before_thinking:
                            yield f"data: {json.dumps({'delta': before_thinking})}\n\n"
                        
                        # Start capturing thinking content
                        in_thinking = True
                        content_buffer = content_buffer[think_start.end():]
                        thinking_buffer = ""
                    else:
                        # No thinking tag found, send accumulated content
                        if content_buffer:
                            yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
                        content_buffer = ""
                        break
                else:
                    # We're inside a thinking tag, look for closing tag
                    think_end = re.search(r'</think(?:ing)?>', content_buffer, re.IGNORECASE)
                    if think_end:
                        # Capture thinking content up to the closing tag
                        thinking_buffer += content_buffer[:think_end.start()]
                        
                        # Send the thinking content as a special type
                        if thinking_buffer.strip():
                            yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
                        
                        # Continue processing after the closing tag
                        in_thinking = False
                        content_buffer = content_buffer[think_end.end():]
                        thinking_buffer = ""
                    else:
                        # Still inside thinking tag, accumulate content
                        thinking_buffer += content_buffer
                        content_buffer = ""
                        break
    
    # Handle any remaining content
    if in_thinking and thinking_buffer:
        # Unclosed thinking tag - send as thinking content
        yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
    elif content_buffer:
        # Regular content
        yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
    
    # Signal the end of the stream
    yield f"data: {json.dumps({'end_of_stream': True})}\n\n"


def format_api_error_message(error_str):
    """
    Formats API error messages to be more user-friendly.
    Specifically handles token limit errors with helpful suggestions.
    """
    error_lower = error_str.lower()
    
    # Check for token limit errors
    if 'maximum context length' in error_lower and 'tokens' in error_lower:
        return "[Summary generation failed: The transcription is too long for AI processing. Request your admin to try using a different LLM with a larger context size, or set a limit for the transcript_length_limit in the system settings.]"
    
    # Check for other common API errors
    if 'rate limit' in error_lower:
        return "[Summary generation failed: API rate limit exceeded. Please try again in a few minutes.]"
    
    if 'insufficient funds' in error_lower or 'quota exceeded' in error_lower:
        return "[Summary generation failed: API quota exceeded. Please contact support.]"
    
    if 'timeout' in error_lower:
        return "[Summary generation failed: Request timed out. Please try again.]"
    
    # For other errors, show a generic message
    return f"[Summary generation failed: {error_str}]"


# --- Document Processing Functions ---

def process_markdown_to_docx(doc, content):
    """Convert markdown content to properly formatted Word document elements.

    Supports:
    - Tables (markdown pipe tables)
    - Headings (# ## ###)
    - Bold text (**text**)
    - Italic text (*text* or _text_)
    - Bold italic (***text***)
    - Inline code (`code`)
    - Code blocks (```code```)
    - Strikethrough (~~text~~)
    - Links ([text](url))
    - Bullet lists (- or *)
    - Numbered lists (1. 2. 3.)
    - Horizontal rules (--- or ***)
    """
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    def ensure_unicode_font(run, text):
        """Ensure the run uses a font that supports the characters in the text."""
        # Check if text contains non-ASCII characters
        try:
            text.encode('ascii')
            # Text is pure ASCII, no special font needed
        except UnicodeEncodeError:
            # Text contains non-ASCII characters, use a font with better Unicode support
            # Use Arial for broad compatibility - it has good Unicode support on most systems
            run.font.name = 'Arial'
            # Set the East Asian font for CJK (Chinese, Japanese, Korean) text
            # This ensures proper rendering in Word
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        return run

    def add_formatted_run(paragraph, text):
        """Add a run with inline formatting to a paragraph."""
        if not text:
            return

        # Pattern for all inline formatting
        # Order matters: check triple asterisk before double/single
        patterns = [
            (r'\*\*\*(.*?)\*\*\*', lambda p, t: (lambda r: (setattr(r, 'bold', True), setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Bold italic
            (r'\*\*(.*?)\*\*', lambda p, t: (lambda r: (setattr(r, 'bold', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Bold
            (r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', lambda p, t: (lambda r: (setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Italic with *
            (r'\b_(.*?)_\b', lambda p, t: (lambda r: (setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Italic with _
            (r'~~(.*?)~~', lambda p, t: (lambda r: (setattr(r, 'strike', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Strikethrough
            (r'`([^`]+)`', lambda p, t: add_code_run(p, t)),  # Inline code
            (r'\[([^\]]+)\]\(([^)]+)\)', lambda p, t, u: add_link_run(p, t, u)),  # Links
        ]

        def add_code_run(para, text):
            """Add inline code with monospace font and background."""
            run = para.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson color for code
            # Check if we need Unicode support for code
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                # Use Consolas as fallback for better Unicode support in monospace
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            return run

        def add_link_run(para, text, url):
            """Add a hyperlink-styled run (note: actual hyperlinks require more complex handling)."""
            full_text = f"{text} ({url})"
            run = para.add_run(full_text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
            run.font.underline = True
            ensure_unicode_font(run, full_text)
            return run

        # Process the text with all patterns
        remaining_text = text
        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            matched_pattern = None

            # Find the earliest matching pattern
            for pattern, handler in patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    matched_pattern = handler

            if earliest_match:
                # Add text before the match
                if earliest_pos > 0:
                    run = paragraph.add_run(remaining_text[:earliest_pos])
                    ensure_unicode_font(run, remaining_text[:earliest_pos])

                # Apply formatting for the matched text
                if '[' in earliest_match.group(0) and '](' in earliest_match.group(0):
                    # Special handling for links (two groups)
                    matched_pattern(paragraph, earliest_match.group(1), earliest_match.group(2))
                else:
                    matched_pattern(paragraph, earliest_match.group(1))

                # Continue with remaining text
                remaining_text = remaining_text[earliest_match.end():]
            else:
                # No more patterns, add the rest as plain text
                run = paragraph.add_run(remaining_text)
                ensure_unicode_font(run, remaining_text)
                break

    def parse_table(lines, start_idx):
        """Parse a markdown table starting at the given index."""
        if start_idx >= len(lines):
            return None, start_idx

        # Check if this looks like a table
        if '|' not in lines[start_idx]:
            return None, start_idx

        table_data = []
        idx = start_idx

        while idx < len(lines) and '|' in lines[idx]:
            # Skip separator lines
            if re.match(r'^[\s\|\-:]+$', lines[idx]):
                idx += 1
                continue

            # Parse cells
            cells = [cell.strip() for cell in lines[idx].split('|')]
            # Remove empty cells at start and end
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]

            if cells:
                table_data.append(cells)
            idx += 1

        if table_data:
            return table_data, idx
        return None, start_idx

    # Split content into lines
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                # End of code block - add it as preformatted text
                in_code_block = False
                if code_block_content:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    code_text = '\n'.join(code_block_content)
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)
                    # Check if we need Unicode support for code blocks
                    try:
                        code_text.encode('ascii')
                    except UnicodeEncodeError:
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Check for table
        table_data, end_idx = parse_table(lines, i)
        if table_data:
            # Create Word table
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'

            # Populate table
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        # Clear existing paragraphs and add new one
                        cell.text = ""
                        p = cell.add_paragraph()
                        add_formatted_run(p, cell_text)
                        # Make header row bold
                        if row_idx == 0:
                            for run in p.runs:
                                run.bold = True

            doc.add_paragraph('')  # Space after table
            i = end_idx
            continue

        line = line.rstrip()

        # Skip empty lines
        if not line:
            doc.add_paragraph('')
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^(\*{3,}|-{3,}|_{3,})$', line.strip()):
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        # Bullet points
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            # Get the indentation level
            indent = len(line) - len(line.lstrip())
            bullet_text = line.lstrip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Add indentation if nested
            if indent > 0:
                p.paragraph_format.left_indent = Pt(indent * 10)
            add_formatted_run(p, bullet_text)
        # Numbered lists
        elif re.match(r'^\s*\d+\.', line):
            match = re.match(r'^(\s*)(\d+)\.\s*(.*)', line)
            if match:
                indent = len(match.group(1))
                list_text = match.group(3)
                p = doc.add_paragraph(style='List Number')
                if indent > 0:
                    p.paragraph_format.left_indent = Pt(indent * 10)
                add_formatted_run(p, list_text)
        # Blockquote
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(30)
            add_formatted_run(p, line[2:])
            # Add a gray color to indicate quote
            for run in p.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            add_formatted_run(p, line)

        i += 1


# --- iCalendar Functions ---

def escape_ical_text(text):
    """Escape special characters for iCalendar format."""
    if not text:
        return ''
    # Escape special characters
    text = str(text)
    text = text.replace('\\', '\\\\')
    text = text.replace(',', '\\,')
    text = text.replace(';', '\\;')
    text = text.replace('\n', '\\n')
    return text


def generate_ics_content(event):
    """Generate ICS calendar file content for an event."""
    
    # Generate unique ID for the event
    uid = f"{event.id}-{uuid.uuid4()}@lecapp.app"

    # Format dates in iCalendar format (YYYYMMDDTHHMMSS)
    def format_ical_date(dt):
        if dt:
            return dt.strftime('%Y%m%dT%H%M%S')
        return None

    # Start building ICS content
    lines = [
        'BEGIN:VCALENDAR',
        'VERSION:2.0',
        'PRODID:-//LecApp//Event Export//EN',
        'CALSCALE:GREGORIAN',
        'METHOD:PUBLISH',
        'BEGIN:VEVENT',
        f'UID:{uid}',
        f'DTSTAMP:{format_ical_date(datetime.utcnow())}',
    ]

    # Add event details
    if event.start_datetime:
        lines.append(f'DTSTART:{format_ical_date(event.start_datetime)}')

    if event.end_datetime:
        lines.append(f'DTEND:{format_ical_date(event.end_datetime)}')
    elif event.start_datetime:
        # If no end time, default to 1 hour after start
        end_time = event.start_datetime + timedelta(hours=1)
        lines.append(f'DTEND:{format_ical_date(end_time)}')

    # Add title and description
    lines.append(f'SUMMARY:{escape_ical_text(event.title)}')

    if event.description:
        lines.append(f'DESCRIPTION:{escape_ical_text(event.description)}')

    # Add location if available
    if event.location:
        lines.append(f'LOCATION:{escape_ical_text(event.location)}')

    # Add attendees if available
    if event.attendees:
        try:
            attendees_list = json.loads(event.attendees)
            for attendee in attendees_list:
                if attendee:
                    lines.append(f'ATTENDEE:CN={escape_ical_text(attendee)}:mailto:{attendee.replace(" ", ".").lower()}@example.com')
        except:
            pass

    # Add reminder/alarm if specified
    if event.reminder_minutes and event.reminder_minutes > 0:
        lines.extend([
            'BEGIN:VALARM',
            'TRIGGER:-PT{}M'.format(event.reminder_minutes),
            'ACTION:DISPLAY',
            f'DESCRIPTION:Reminder: {escape_ical_text(event.title)}',
            'END:VALARM'
        ])

    # Close event and calendar
    lines.extend([
        'STATUS:CONFIRMED',
        'TRANSP:OPAQUE',
        'END:VEVENT',
        'END:VCALENDAR'
    ])

    return '\r\n'.join(lines)