"""Recordings blueprint handling uploads, status, toggles, and basic recording operations.

Incremental extraction (Task 9). Additional routes (downloads, speakers, events, chat, tags, shares) will follow.
"""
import os
import mimetypes
import subprocess
import threading
from datetime import datetime, timedelta
from sqlalchemy import select

from flask import Blueprint, request, jsonify, make_response, send_file, Response, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

from src.extensions import db, limiter
from src.models import Recording, Tag, RecordingTag, Event, Speaker, TranscriptTemplate, SystemSetting, Share, Attachment, TranscriptChunk
from src.services.transcription_service import (
    transcribe_audio_task, USE_ASR_ENDPOINT, ENABLE_CHUNKING, chunking_service,
    ASR_MIN_SPEAKERS, ASR_MAX_SPEAKERS
)
from src.services.llm_service import (
    call_llm_completion, generate_summary_only_task, identify_unidentified_speakers_from_text
)
from src.utils import (
    process_markdown_to_docx, format_transcription_for_llm, process_streaming_with_thinking,
    generate_ics_content, sanitize_html
)
import json
import re

recordings_bp = Blueprint('recordings', __name__)


@recordings_bp.route('/api/recordings', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")
def get_recordings_paginated():
    """Get recordings with pagination and server-side filtering."""
    try:
        # Parse query parameters
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 25, type=int), 100)  # Cap at 100 per page
        search_query = request.args.get('q', '').strip()
        
        # Build base query
        stmt = select(Recording).where(Recording.user_id == current_user.id)
        
        # Apply search filters if provided
        if search_query:
            # Parse search query for special syntax
            
            # Extract date filters
            date_filters = re.findall(r'date:(\S+)', search_query.lower())
            date_from_filters = re.findall(r'date_from:(\S+)', search_query.lower())
            date_to_filters = re.findall(r'date_to:(\S+)', search_query.lower())
            tag_filters = re.findall(r'tag:(\S+)', search_query.lower())
            
            # Remove special syntax to get text search
            text_query = re.sub(r'date:\S+', '', search_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_from:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_to:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'tag:\S+', '', text_query, flags=re.IGNORECASE).strip()
            
            # Apply date filters
            for date_filter in date_filters:
                if date_filter == 'today':
                    today = datetime.now().date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == today,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == today
                            )
                        )
                    )
                elif date_filter == 'yesterday':
                    yesterday = datetime.now().date() - timedelta(days=1)
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == yesterday,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == yesterday
                            )
                        )
                    )
                elif date_filter == 'thisweek':
                    today = datetime.now().date()
                    start_of_week = today - timedelta(days=today.weekday())
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_week,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_week
                            )
                        )
                    )
                elif date_filter == 'lastweek':
                    today = datetime.now().date()
                    end_of_last_week = today - timedelta(days=today.weekday())
                    start_of_last_week = end_of_last_week - timedelta(days=7)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= start_of_last_week,
                                Recording.meeting_date < end_of_last_week
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_last_week,
                                db.func.date(Recording.created_at) < end_of_last_week
                            )
                        )
                    )
                elif date_filter == 'thismonth':
                    today = datetime.now().date()
                    start_of_month = today.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_month,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_month
                            )
                        )
                    )
                elif date_filter == 'lastmonth':
                    today = datetime.now().date()
                    first_day_this_month = today.replace(day=1)
                    last_day_last_month = first_day_this_month - timedelta(days=1)
                    first_day_last_month = last_day_last_month.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= first_day_last_month,
                                Recording.meeting_date <= last_day_last_month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= first_day_last_month,
                                db.func.date(Recording.created_at) <= last_day_last_month
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', date_filter):
                    # Specific date format YYYY-MM-DD
                    target_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == target_date,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == target_date
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}$', date_filter):
                    # Month format YYYY-MM
                    year, month = map(int, date_filter.split('-'))
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                db.extract('year', Recording.meeting_date) == year,
                                db.extract('month', Recording.meeting_date) == month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year,
                                db.extract('month', Recording.created_at) == month
                            )
                        )
                    )
                elif re.match(r'^\d{4}$', date_filter):
                    # Year format YYYY
                    year = int(date_filter)
                    stmt = stmt.where(
                        db.or_(
                            db.extract('year', Recording.meeting_date) == year,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year
                            )
                        )
                    )
            
            # Apply date range filters
            if date_from_filters and date_from_filters[0]:
                try:
                    date_from = datetime.strptime(date_from_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= date_from,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= date_from
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore
            
            if date_to_filters and date_to_filters[0]:
                try:
                    date_to = datetime.strptime(date_to_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date <= date_to,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) <= date_to
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore
            
            # Apply tag filters
            if tag_filters:
                # Join with tags table and filter by tag names
                tag_conditions = []
                for tag_filter in tag_filters:
                    # Replace underscores back to spaces for matching
                    tag_name = tag_filter.replace('_', ' ')
                    tag_conditions.append(Tag.name.ilike(f'%{tag_name}%'))
                
                stmt = stmt.join(RecordingTag).join(Tag).where(db.or_(*tag_conditions))
            
            # Apply text search
            if text_query:
                text_conditions = [
                    Recording.title.ilike(f'%{text_query}%'),
                    Recording.participants.ilike(f'%{text_query}%'),
                    Recording.transcription.ilike(f'%{text_query}%'),
                    Recording.notes.ilike(f'%{text_query}%')
                ]
                stmt = stmt.where(db.or_(*text_conditions))
        
        # Apply ordering (most recent first based on meeting_date or created_at)
        stmt = stmt.order_by(
            db.case(
                (Recording.meeting_date.is_not(None), Recording.meeting_date),
                else_=db.func.date(Recording.created_at)
            ).desc(),
            Recording.created_at.desc()
        )
        
        # Get total count for pagination info
        count_stmt = select(db.func.count()).select_from(stmt.subquery())
        total_count = db.session.execute(count_stmt).scalar()
        
        # Apply pagination
        offset = (page - 1) * per_page
        stmt = stmt.offset(offset).limit(per_page)
        
        # Execute query
        recordings = db.session.execute(stmt).scalars().all()
        
        # Calculate pagination metadata
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_prev = page > 1
        
        return jsonify({
            'recordings': [recording.to_dict() for recording in recordings],
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_prev': has_prev
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Error fetching paginated recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/inbox_recordings', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")
def get_inbox_recordings():
    """Get recordings that are in the inbox and currently processing."""
    try:
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.is_inbox == True,
            Recording.status.in_(['PENDING', 'PROCESSING', 'SUMMARIZING'])
        ).order_by(Recording.created_at.desc())
        
        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_dict() for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching inbox recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/recording/<int:recording_id>/toggle_inbox', methods=['POST'])
@login_required
def toggle_inbox(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403
        recording.is_inbox = not recording.is_inbox
        db.session.commit()
        return jsonify({'success': True, 'is_inbox': recording.is_inbox})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling inbox status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/recording/<int:recording_id>/toggle_highlight', methods=['POST'])
@login_required
def toggle_highlight(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403
        recording.is_highlighted = not recording.is_highlighted
        db.session.commit()
        return jsonify({'success': True, 'is_highlighted': recording.is_highlighted})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling highlighted status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        upload_root = request.app.config['UPLOAD_FOLDER'] if hasattr(request, 'app') else None  # fallback
        from flask import current_app as app
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_filename}")
        file.seek(0, os.SEEK_END)
        original_file_size = file.tell()
        file.seek(0)
        max_content_length = app.config.get('MAX_CONTENT_LENGTH')
        should_enforce_size_limit = True
        if ENABLE_CHUNKING and chunking_service and not USE_ASR_ENDPOINT:
            should_enforce_size_limit = False
            mode, limit_value = chunking_service.parse_chunk_limit()
            if mode == 'size':
                app.logger.info(f"Size-based chunking enabled ({limit_value}MB limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
            else:
                app.logger.info(f"Duration-based chunking enabled ({limit_value}s limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
        if should_enforce_size_limit and max_content_length and original_file_size > max_content_length:
            raise RequestEntityTooLarge()
        file.save(filepath)
        app.logger.info(f"File saved to {filepath}")
        filename_lower = original_filename.lower()
        needs_chunking_for_processing = (chunking_service and ENABLE_CHUNKING and not USE_ASR_ENDPOINT and chunking_service.needs_chunking(filepath, USE_ASR_ENDPOINT))
        if needs_chunking_for_processing:
            supported_formats = ('.wav', '.mp3', '.flac')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.m4a', '.aac', '.ogg', '.wma', '.webm')
        else:
            supported_formats = ('.wav', '.mp3', '.flac', '.webm', '.m4a', '.aac', '.ogg')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.wma')
        is_problematic_aac = (USE_ASR_ENDPOINT and (filename_lower.endswith('.aac') or 'aac' in filename_lower))
        should_convert = ((not filename_lower.endswith(supported_formats) and needs_chunking_for_processing) or is_problematic_aac)
        if should_convert:
            base_filepath, _ = os.path.splitext(filepath)
            temp_mp3_filepath = f"{base_filepath}_temp.mp3"
            mp3_filepath = f"{base_filepath}.mp3"
            try:
                subprocess.run(
                    ['ffmpeg', '-i', filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', '-ac', '1', temp_mp3_filepath],
                    check=True, capture_output=True, text=True
                )
                if filepath.lower() != mp3_filepath.lower():
                    os.remove(filepath)
                os.rename(temp_mp3_filepath, mp3_filepath)
                filepath = mp3_filepath
            except Exception as e:
                app.logger.error(f"Audio conversion failed: {e}")
                return jsonify({'error': 'Failed to convert audio file.'}), 500
        final_file_size = os.path.getsize(filepath)
        mime_type, _ = mimetypes.guess_type(filepath)
        notes = request.form.get('notes')
        selected_tags = []
        tag_index = 0
        while True:
            tag_id_key = f'tag_ids[{tag_index}]'
            tag_id = request.form.get(tag_id_key)
            if not tag_id:
                break
            tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first()
            if tag:
                selected_tags.append(tag)
            tag_index += 1
        if not selected_tags:
            single_tag_id = request.form.get('tag_id')
            if single_tag_id:
                tag = Tag.query.filter_by(id=single_tag_id, user_id=current_user.id).first()
                if tag:
                    selected_tags.append(tag)
        language = request.form.get('language', '')
        min_speakers = request.form.get('min_speakers') or None
        max_speakers = request.form.get('max_speakers') or None
        for var_name in ['min_speakers', 'max_speakers']:
            val = locals()[var_name]
            if val:
                try:
                    locals()[var_name] = int(val)
                except (ValueError, TypeError):
                    locals()[var_name] = None
        if selected_tags:
            first_tag = selected_tags[0]
            if not language and first_tag.default_language:
                language = first_tag.default_language
            if min_speakers is None and first_tag.default_min_speakers:
                min_speakers = first_tag.default_min_speakers
            if max_speakers is None and first_tag.default_max_speakers:
                max_speakers = first_tag.default_max_speakers
        if min_speakers is None and ASR_MIN_SPEAKERS:
            try:
                min_speakers = int(ASR_MIN_SPEAKERS)
            except (ValueError, TypeError):
                pass
        if max_speakers is None and ASR_MAX_SPEAKERS:
            try:
                max_speakers = int(ASR_MAX_SPEAKERS)
            except (ValueError, TypeError):
                pass
        now = datetime.utcnow()
        recording = Recording(
            audio_path=filepath,
            original_filename=original_filename,
            title=f"Recording - {original_filename}",
            file_size=final_file_size,
            status='PENDING',
            meeting_date=now.date(),
            user_id=current_user.id,
            mime_type=mime_type,
            notes=notes,
            processing_source='upload'
        )
        db.session.add(recording)
        db.session.commit()
        for order, tag in enumerate(selected_tags, 1):
            db.session.add(RecordingTag(recording_id=recording.id, tag_id=tag.id, order=order, added_at=datetime.utcnow()))
        if selected_tags:
            db.session.commit()
        app.logger.info(f"Initial recording record created with ID: {recording.id}")
        start_time = datetime.utcnow()
        first_tag = selected_tags[0] if selected_tags else None
        if USE_ASR_ENDPOINT:
            thread = threading.Thread(
                target=transcribe_audio_task,
                args=(current_app._get_current_object().app_context(), recording.id, filepath, os.path.basename(filepath), start_time),
                kwargs={'language': language, 'min_speakers': min_speakers, 'max_speakers': max_speakers, 'tag_id': first_tag.id if first_tag else None}
            )
        else:
            thread = threading.Thread(
                target=transcribe_audio_task,
                args=(current_app._get_current_object().app_context(), recording.id, filepath, os.path.basename(filepath), start_time),
                kwargs={'tag_id': first_tag.id if first_tag else None}
            )
        thread.start()
        response_data = recording.to_dict()
        response_data['recording_id'] = recording.id  # For backward compatibility with tests
        return jsonify(response_data), 202
    except RequestEntityTooLarge:
        from flask import current_app as app
        max_size_mb = app.config['MAX_CONTENT_LENGTH'] / (1024 * 1024)
        return jsonify({'error': f'File too large. Maximum size is {max_size_mb:.0f} MB.'}), 413
    except Exception as e:
        from flask import current_app as app
        db.session.rollback()
        app.logger.error(f"Error during file upload: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred during upload.'}), 500


@recordings_bp.route('/status/<int:recording_id>', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")
def get_status(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to view this recording'}), 403
        db.session.refresh(recording)
        return jsonify(recording.to_dict())
    except Exception as e:
        current_app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# --- Speaker Management Routes ---
@recordings_bp.route('/speakers', methods=['GET'])
@login_required
def get_speakers():
    try:
        speakers = Speaker.query.filter_by(user_id=current_user.id) \
            .order_by(Speaker.use_count.desc(), Speaker.last_used.desc()).all()
        return jsonify([s.to_dict() for s in speakers])
    except Exception as e:
        current_app.logger.error(f"Error fetching speakers: {e}")
        return jsonify({'error': str(e)}), 500

@recordings_bp.route('/speakers/search', methods=['GET'])
@login_required
def search_speakers():
    try:
        query = request.args.get('q', '').strip()
        if not query:
            return jsonify([])
        speakers = Speaker.query.filter_by(user_id=current_user.id) \
            .filter(Speaker.name.ilike(f'%{query}%')) \
            .order_by(Speaker.use_count.desc(), Speaker.last_used.desc()) \
            .limit(10).all()
        return jsonify([s.to_dict() for s in speakers])
    except Exception as e:
        current_app.logger.error(f"Error searching speakers: {e}")
        return jsonify({'error': str(e)}), 500

@recordings_bp.route('/speakers', methods=['POST'])
@login_required
def create_speaker():
    try:
        data = request.json or {}
        name = data.get('name', '').strip()
        if not name:
            return jsonify({'error': 'Speaker name is required'}), 400
        existing = Speaker.query.filter_by(user_id=current_user.id, name=name).first()
        if existing:
            existing.use_count += 1
            existing.last_used = datetime.utcnow()
            db.session.commit()
            return jsonify(existing.to_dict())
        speaker = Speaker(name=name, user_id=current_user.id, use_count=1, created_at=datetime.utcnow(), last_used=datetime.utcnow())
        db.session.add(speaker)
        db.session.commit()
        return jsonify(speaker.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating speaker: {e}")
        return jsonify({'error': str(e)}), 500

@recordings_bp.route('/speakers/<int:speaker_id>', methods=['DELETE'])
@login_required
def delete_speaker(speaker_id):
    try:
        speaker = Speaker.query.filter_by(id=speaker_id, user_id=current_user.id).first()
        if not speaker:
            return jsonify({'error': 'Speaker not found'}), 404
        db.session.delete(speaker)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting speaker: {e}")
        return jsonify({'error': str(e)}), 500

@recordings_bp.route('/speakers/delete_all', methods=['DELETE'])
@login_required
def delete_all_speakers():
    try:
        deleted = Speaker.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({'success': True, 'deleted_count': deleted})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting all speakers: {e}")
        return jsonify({'error': str(e)}), 500


# --- Downloads (transcript, summary, notes, chat) ---
def _add_unicode_paragraph(doc, text):
    p = doc.add_paragraph(text)
    try:
        text.encode('ascii')
    except UnicodeEncodeError:
        from docx.oxml.ns import qn
        for run in p.runs:
            run.font.name = 'Arial'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return p

@recordings_bp.route('/recording/<int:recording_id>/download/transcript')
@login_required
def download_transcript_with_template(recording_id):
    try:
        from datetime import timedelta
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this recording'}), 403
        if not recording.transcription:
            return jsonify({'error': 'No transcription available for this recording'}), 400
        template_id = request.args.get('template_id', type=int)
        if template_id:
            template = TranscriptTemplate.query.filter_by(id=template_id, user_id=current_user.id).first()
        else:
            template = TranscriptTemplate.query.filter_by(user_id=current_user.id, is_default=True).first()
        template_format = template.template if template else "[{{speaker}}]: {{text}}"
        def format_time(seconds):
            if seconds is None: return "00:00:00"
            td = timedelta(seconds=seconds); h=int(td.total_seconds()//3600); m=int((td.total_seconds()%3600)//60); s=int(td.total_seconds()%60); return f"{h:02d}:{m:02d}:{s:02d}"
        def format_srt_time(seconds):
            if seconds is None: return "00:00:00,000"
            td = timedelta(seconds=seconds); h=int(td.total_seconds()//3600); m=int((td.total_seconds()%3600)//60); s=int(td.total_seconds()%60); ms=int((td.total_seconds()%1)*1000); return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"
        try:
            transcription_data = json.loads(recording.transcription)
        except Exception:
            return jsonify({'error': 'Invalid transcription format'}), 400
        output_lines=[]
        for index, segment in enumerate(transcription_data,1):
            line = template_format
            replacements={'{{index}}':str(index),'{{speaker}}':segment.get('speaker','Unknown'),'{{text}}':segment.get('sentence',''),'{{start_time}}':format_time(segment.get('start_time')),'{{end_time}}':format_time(segment.get('end_time'))}
            for k,v in replacements.items(): line=line.replace(k,v)
            line = re.sub(r'{{(.*?)\|upper}}', lambda m: replacements.get('{{'+m.group(1)+'}}','').upper(), line)
            line = re.sub(r'{{start_time\|srt}}', format_srt_time(segment.get('start_time')), line)
            line = re.sub(r'{{end_time\|srt}}', format_srt_time(segment.get('end_time')), line)
            output_lines.append(line)
        formatted='\n'.join(output_lines)
        response = make_response(formatted)
        filename = f"{(recording.title or 'transcript')}_{(template.name if template else 'formatted')}.txt"
        filename = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
        response.headers['Content-Type']='text/plain; charset=utf-8'
        response.headers['Content-Disposition']=f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        current_app.logger.error(f"Error downloading transcript: {e}")
        return jsonify({'error': 'Failed to generate transcript download'}), 500

def _set_rfc2231_filename(response, filename, fallback):
    try:
        filename.encode('ascii')
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    except UnicodeEncodeError:
        from urllib.parse import quote
        try:
            encoded = quote(filename)
            response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded}"
        except Exception:
            response.headers['Content-Disposition'] = f'attachment; filename="{fallback}"'
    return response

def _word_doc_filename(prefix, recording, recording_id, base):
    safe_title = re.sub(r'[<>:"/\\|?*]','', recording.title or 'Untitled')
    safe_title = re.sub(r'[-\s]+','-', safe_title).strip('-')
    filename = f'{prefix}-{safe_title}.docx' if safe_title else f'{prefix}-recording-{recording_id}.docx'
    ascii_filename = filename.encode('ascii','ignore').decode('ascii')
    if not ascii_filename.strip() or ascii_filename.strip() in [f'{prefix}-.docx', f'{prefix}-recording-.docx']:
        ascii_filename = f'{prefix}-recording-{recording_id}.docx'
    return filename, ascii_filename

def _apply_unicode_heading(run_container, text):
    try: text.encode('ascii')
    except UnicodeEncodeError:
        from docx.oxml.ns import qn
        for run in run_container.runs:
            run.font.name='Arial'; r=run._element; r.rPr.rFonts.set(qn('w:eastAsia'),'Arial')

@recordings_bp.route('/recording/<int:recording_id>/download/summary')
@login_required
def download_summary_word(recording_id):
    try:
        from docx import Document
        from io import BytesIO
        recording = db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to access this recording'}),403
        if not recording.summary: return jsonify({'error':'No summary available for this recording'}),400
        doc=Document(); title_text=f'Summary: {recording.title or "Untitled Recording"}'; title=doc.add_heading(title_text,0); _apply_unicode_heading(title, title_text)
        _add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date: _add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants: _add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        if recording.tags: _add_unicode_paragraph(doc, f'Tags: {", ".join([t.name for t in recording.tags])}')
        doc.add_paragraph('')
        process_markdown_to_docx(doc, recording.summary)
        buf=BytesIO(); doc.save(buf); buf.seek(0)
        filename, ascii_filename = _word_doc_filename('summary', recording, recording_id, recording.title)
        response = send_file(buf, as_attachment=False, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return _set_rfc2231_filename(response, filename, f'summary-recording-{recording_id}.docx')
    except Exception as e:
        current_app.logger.error(f"Error generating summary Word document: {e}")
        return jsonify({'error':'Failed to generate Word document'}),500

@recordings_bp.route('/recording/<int:recording_id>/download/chat', methods=['POST'])
@login_required
def download_chat_word(recording_id):
    try:
        from docx import Document
        from io import BytesIO
        recording=db.session.get(Recording,recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to access this recording'}),403
        data=request.json or {}
        messages=data.get('messages')
        if not messages: return jsonify({'error':'No messages to download'}),400
        doc=Document(); title_text=f'Chat Conversation: {recording.title or "Untitled Recording"}'; title=doc.add_heading(title_text,0); _apply_unicode_heading(title,title_text)
        _add_unicode_paragraph(doc, f'Recording Date: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        _add_unicode_paragraph(doc, f'Chat Export Date: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph('')
        for message in messages:
            role = message.get('role','unknown'); content=message.get('content',''); thinking=message.get('thinking','')
            p=doc.add_paragraph(); run=p.add_run(f'{"You" if role=="user" else ("Assistant" if role=="assistant" else role.title())}: '); run.bold=True
            if thinking and role=='assistant':
                p=doc.add_paragraph(); p.add_run('[Model Reasoning]\n').italic=True; p.add_run(thinking).italic=True; doc.add_paragraph('')
            process_markdown_to_docx(doc, content); doc.add_paragraph('')
        buf=BytesIO(); doc.save(buf); buf.seek(0)
        filename, ascii_filename = _word_doc_filename('chat', recording, recording_id, recording.title)
        response=send_file(buf, as_attachment=False, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return _set_rfc2231_filename(response, filename, f'chat-recording-{recording_id}.docx')
    except Exception as e:
        current_app.logger.error(f"Error generating chat Word document: {e}")
        return jsonify({'error':'Failed to generate Word document'}),500

@recordings_bp.route('/recording/<int:recording_id>/download/notes')
@login_required
def download_notes_word(recording_id):
    try:
        from docx import Document
        from io import BytesIO
        recording=db.session.get(Recording,recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to access this recording'}),403
        if not recording.notes: return jsonify({'error':'No notes available for this recording'}),400
        doc=Document(); title_text=f'Notes: {recording.title or "Untitled Recording"}'; title=doc.add_heading(title_text,0); _apply_unicode_heading(title,title_text)
        _add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date: _add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants: _add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        if recording.tags: _add_unicode_paragraph(doc, f'Tags: {", ".join([t.name for t in recording.tags])}')
        doc.add_paragraph(''); process_markdown_to_docx(doc, recording.notes)
        buf=BytesIO(); doc.save(buf); buf.seek(0)
        filename, ascii_filename = _word_doc_filename('notes', recording, recording_id, recording.title)
        response=send_file(buf, as_attachment=False, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return _set_rfc2231_filename(response, filename, f'notes-recording-{recording_id}.docx')
    except Exception as e:
        current_app.logger.error(f"Error generating notes Word document: {e}")
        return jsonify({'error':'Failed to generate Word document'}),500


# --- Events & ICS ---
@recordings_bp.route('/api/recording/<int:recording_id>/events', methods=['GET'])
@login_required
def get_recording_events(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!= current_user.id: return jsonify({'error':'Unauthorized'}),403
        events=Event.query.filter_by(recording_id=recording_id).all()
        return jsonify({'events':[e.to_dict() for e in events]})
    except Exception as e:
        current_app.logger.error(f"Error fetching events for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}),500

@recordings_bp.route('/api/event/<int:event_id>/ics', methods=['GET'])
@login_required
def download_event_ics(event_id):
    try:
        event=db.session.get(Event, event_id)
        if not event: return jsonify({'error':'Event not found'}),404
        if event.recording.user_id != current_user.id: return jsonify({'error':'Unauthorized'}),403
        ics_content=generate_ics_content(event)
        response=make_response(ics_content)
        response.headers['Content-Type']='text/calendar; charset=utf-8'
        response.headers['Content-Disposition']=f'attachment; filename="{secure_filename(event.title)}.ics"'
        return response
    except Exception as e:
        current_app.logger.error(f"Error generating ICS for event {event_id}: {e}")
        return jsonify({'error': str(e)}),500

@recordings_bp.route('/api/recording/<int:recording_id>/events/ics', methods=['GET'])
@login_required
def download_all_events_ics(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id != current_user.id: return jsonify({'error':'Unauthorized'}),403
        events=Event.query.filter_by(recording_id=recording_id).all()
        if not events: return jsonify({'error':'No events found for this recording'}),404
        lines=["BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//LecApp//Event Export//EN","CALSCALE:GREGORIAN","METHOD:PUBLISH"]
        for event in events:
            individual=generate_ics_content(event).split('\n')
            in_event=False
            for line in individual:
                if line.startswith('BEGIN:VEVENT'): in_event=True
                if in_event: lines.append(line)
                if line.startswith('END:VEVENT'): in_event=False
        lines.append("END:VCALENDAR")
        response=make_response('\n'.join(lines))
        response.headers['Content-Type']='text/calendar; charset=utf-8'
        filename=f"events-{secure_filename(recording.title or f'recording-{recording_id}')}.ics"
        response.headers['Content-Disposition']=f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        current_app.logger.error(f"Error generating ICS for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}),500


# --- Metadata & Transcription Update ---
@recordings_bp.route('/save', methods=['POST'])
@login_required
def save_metadata():
    try:
        data=request.json
        if not data: return jsonify({'error':'No data provided'}),400
        recording_id=data.get('id')
        if not recording_id: return jsonify({'error':'No recording ID provided'}),400
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to edit this recording'}),403
        if 'title' in data: recording.title=data['title']
        if 'participants' in data: recording.participants=data['participants']
        if 'notes' in data: recording.notes=sanitize_html(data['notes']) if data['notes'] else data['notes']
        if 'summary' in data: recording.summary=sanitize_html(data['summary']) if data['summary'] else data['summary']
        if 'is_inbox' in data: recording.is_inbox=data['is_inbox']
        if 'is_highlighted' in data: recording.is_highlighted=data['is_highlighted']
        if 'meeting_date' in data:
            try:
                date_str=data['meeting_date']; recording.meeting_date=datetime.strptime(date_str,'%Y-%m-%d').date() if date_str else None
            except (ValueError, TypeError):
                current_app.logger.warning(f"Could not parse meeting_date '{data.get('meeting_date')}'")
        db.session.commit(); return jsonify({'success':True,'recording':recording.to_dict()})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error saving metadata for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':'An unexpected error occurred while saving.'}),500

@recordings_bp.route('/recording/<int:recording_id>/update_transcription', methods=['POST'])
@login_required
def update_transcription(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to edit this recording'}),403
        data=request.json or {}; new_transcription=data.get('transcription')
        if new_transcription is None: return jsonify({'error':'No transcription data provided'}),400
        recording.transcription=new_transcription; db.session.commit(); current_app.logger.info(f"Transcription for recording {recording_id} updated")
        return jsonify({'success':True,'message':'Transcription updated successfully.','recording':recording.to_dict()})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error updating transcription for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':'An unexpected error occurred while updating the transcription.'}),500


# --- Summary Generation & Speaker Updates ---
@recordings_bp.route('/recording/<int:recording_id>/generate_summary', methods=['POST'])
@login_required
def generate_summary_endpoint(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to generate summary for this recording'}),403
        if not recording.transcription or len(recording.transcription.strip())<10: return jsonify({'error':'No valid transcription available for summary generation'}),400
        if recording.status in ['PROCESSING','SUMMARIZING']: return jsonify({'error':'Recording is already being processed'}),400
        from flask import current_app as app
        if call_llm_completion is None: return jsonify({'error':'Summary service is not available (OpenRouter client not configured)'}),503
        thread=threading.Thread(target=generate_summary_only_task, args=(app.app_context(), recording_id))
        thread.start(); return jsonify({'success':True,'message':'Summary generation started'})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error starting summary generation for recording {recording_id}: {e}")
        return jsonify({'error':str(e)}),500

def _update_speaker_usage(names):
    if not names or not current_user.is_authenticated: return
    try:
        for name in names:
            name = name.strip()
            if not name:
                continue
            speaker=Speaker.query.filter_by(user_id=current_user.id, name=name).first()
            if speaker:
                speaker.use_count +=1; speaker.last_used=datetime.utcnow()
            else:
                db.session.add(Speaker(name=name,user_id=current_user.id,use_count=1,created_at=datetime.utcnow(),last_used=datetime.utcnow()))
        db.session.commit()
    except Exception as e:
        current_app.logger.error(f"Error updating speaker usage: {e}"); db.session.rollback()

@recordings_bp.route('/recording/<int:recording_id>/update_speakers', methods=['POST'])
@login_required
def update_speakers(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to edit this recording'}),403
        data=request.json or {}; speaker_map=data.get('speaker_map'); regenerate=data.get('regenerate_summary',False)
        if not speaker_map: return jsonify({'error':'No speaker map provided'}),400
        transcription_text=recording.transcription; is_json=False
        try:
            transcription_data=json.loads(transcription_text); is_json=isinstance(transcription_data,list)
        except Exception: is_json=False
        used=[]
        if is_json:
            for seg in transcription_data:
                original=seg.get('speaker');
                if original in speaker_map:
                    new_info=speaker_map[original]; new_name=new_info.get('name','').strip();
                    if new_info.get('isMe'): new_name=current_user.name or 'Me'
                    if new_name:
                        seg['speaker']=new_name
                        if new_name not in used: used.append(new_name)
            recording.transcription=json.dumps(transcription_data)
            final=set()
            for seg in transcription_data:
                sp=seg.get('speaker')
                if sp and str(sp).strip() and not re.match(r'^SPEAKER_\d+$', str(sp), re.IGNORECASE): final.add(sp)
            recording.participants=', '.join(sorted(list(final)))
        else:
            for label, info in speaker_map.items():
                new_name=info.get('name','').strip();
                if info.get('isMe'): new_name=current_user.name or 'Me'
                if new_name:
                    transcription_text=re.sub(r'\[\s*'+re.escape(label)+r'\s*\]', f'[{new_name}]', transcription_text, flags=re.IGNORECASE)
                    if new_name not in used: used.append(new_name)
            recording.transcription=transcription_text
            if used: recording.participants=', '.join(used)
        if used: _update_speaker_usage(used)
        db.session.commit()
        if regenerate:
            from flask import current_app as app
            thread=threading.Thread(target=generate_summary_only_task, args=(app.app_context(), recording.id)); thread.start()
        return jsonify({'success':True,'message':'Speakers updated successfully.','recording':recording.to_dict()})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error updating speakers for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':str(e)}),500

@recordings_bp.route('/recording/<int:recording_id>/auto_identify_speakers', methods=['POST'])
@login_required
def auto_identify_speakers(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to modify this recording'}),403
        if not recording.transcription: return jsonify({'error':'No transcription available for speaker identification'}),400
        data=request.json or {}; current_map=data.get('current_speaker_map', {})
        formatted=format_transcription_for_llm(recording.transcription)
        all_labels=re.findall(r'\[(SPEAKER_\d+)\]', formatted); seen=set(); labels=[x for x in all_labels if not (x in seen or seen.add(x))]
        unidentified=[]
        for label in labels:
            info=current_map.get(label,{}); name=(info.get('name','').strip() if isinstance(info,dict) else str(info).strip())
            if not name: unidentified.append(label)
        if not unidentified: return jsonify({'success':True,'speaker_map':{},'message':'All speakers are already identified'})
        speaker_map=identify_unidentified_speakers_from_text(recording.transcription, unidentified)
        return jsonify({'success':True,'speaker_map':speaker_map})
    except ValueError as ve:
        return jsonify({'error':str(ve)}),503
    except Exception as e:
        current_app.logger.error(f"Error during auto speaker identification for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':f'An unexpected error occurred: {e}'}),500


# --- Chat with Transcription ---
@recordings_bp.route('/chat', methods=['POST'])
@login_required
def chat_with_transcription():
    try:
        data=request.json or {}
        recording_id=data.get('recording_id'); user_message=data.get('message'); history=data.get('message_history', [])
        if not recording_id: return jsonify({'error':'No recording ID provided'}),400
        if not user_message: return jsonify({'error':'No message provided'}),400
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to chat with this recording'}),403
        if call_llm_completion is None: return jsonify({'error':'Chat service is not available (OpenRouter client not configured)'}),503
        user_lang=current_user.output_language if current_user.is_authenticated else None
        instruction=f"Please provide all your responses in {user_lang}." if user_lang else ""
        user_name=current_user.name or 'User'; user_title=current_user.job_title or 'a professional'; user_company=current_user.company or 'their organization'
        formatted=format_transcription_for_llm(recording.transcription)
        transcript_limit=SystemSetting.get_setting('transcript_length_limit', 30000)
        chat_transcript=formatted if transcript_limit==-1 else formatted[:transcript_limit]
        system_prompt=f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {instruction} Analyze the following meeting information and respond to the specific request.

Following are the meeting participants and their roles:
{recording.participants or 'No specific participants information provided.'}

Following is the meeting transcript:
<<start transcript>>
{chat_transcript or 'No transcript available.'}
<<end transcript>>

Additional context and notes about the meeting:
{recording.notes or 'none'}
"""
        messages=[{"role":"system","content":system_prompt}];
        if history: messages.extend(history)
        messages.append({"role":"user","content":user_message})
        def generate():
            try:
                stream=call_llm_completion(messages=messages, temperature=0.7, max_tokens=int(os.environ.get('CHAT_MAX_TOKENS','2000')), stream=True)
                for resp in process_streaming_with_thinking(stream): yield resp
            except Exception as e:
                current_app.logger.error(f"Error during chat stream generation: {e}")
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
        return Response(generate(), mimetype='text/event-stream')
    except Exception as e:
        current_app.logger.error(f"Error in chat endpoint: {e}")
        return jsonify({'error':str(e)}),500


# --- Reprocessing Endpoints ---
@recordings_bp.route('/recording/<int:recording_id>/reprocess_transcription', methods=['POST'])
@login_required
def reprocess_transcription(recording_id):
    try:
        from flask import current_app as app
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to reprocess this recording'}),403
        if not recording.audio_path or not os.path.exists(recording.audio_path): return jsonify({'error':'Audio file not found for reprocessing'}),404
        if recording.status in ['PROCESSING','SUMMARIZING']: return jsonify({'error':'Recording is already being processed'}),400
        filepath=recording.audio_path; filename_for_asr=recording.original_filename or os.path.basename(filepath); lower=filename_for_asr.lower()
        supported=('.wav','.mp3','.flac')
        if not lower.endswith(supported):
            base,ext=os.path.splitext(filepath); temp=f"{base}_temp.mp3"; final=f"{base}.mp3"
            try:
                subprocess.run(['ffmpeg','-i',filepath,'-y','-acodec','libmp3lame','-b:a','128k','-ar','44100',temp],check=True,capture_output=True,text=True)
                if filepath.lower()!=final.lower(): os.remove(filepath)
                os.rename(temp,final); filepath=final; filename_for_asr=os.path.basename(filepath); recording.audio_path=filepath; recording.mime_type, _ = mimetypes.guess_type(filepath); db.session.commit()
            except FileNotFoundError:
                return jsonify({'error':'Audio conversion tool (ffmpeg) not found on server.'}),500
            except subprocess.CalledProcessError as e:
                return jsonify({'error': f'Failed to convert audio file: {e.stderr}'}),500
        recording.transcription=None; recording.summary=None; recording.status='PROCESSING'; Event.query.filter_by(recording_id=recording_id).delete(); db.session.commit(); db.session.refresh(recording)
        if USE_ASR_ENDPOINT:
            data=request.json or {}; language=data.get('language') or (recording.owner.transcription_language if getattr(recording,'owner',None) else None); min_s=data.get('min_speakers') or None; max_s=data.get('max_speakers') or None
            for var in ['min_s','max_s']:
                try:
                    if locals()[var]: locals()[var]=int(locals()[var])
                except (ValueError,TypeError): locals()[var]=None
            if (min_s is None or max_s is None) and recording.tags:
                for assoc in sorted(recording.tag_associations, key=lambda x: x.order):
                    t=assoc.tag
                    if min_s is None and t.default_min_speakers: min_s=t.default_min_speakers
                    if max_s is None and t.default_max_speakers: max_s=t.default_max_speakers
                    if min_s is not None and max_s is not None: break
            if min_s is None and ASR_MIN_SPEAKERS:
                try: min_s=int(ASR_MIN_SPEAKERS)
                except (ValueError,TypeError): pass
            if max_s is None and ASR_MAX_SPEAKERS:
                try: max_s=int(ASR_MAX_SPEAKERS)
                except (ValueError,TypeError): pass
            start_time=datetime.utcnow(); thread=threading.Thread(target=transcribe_audio_task, args=(app.app_context(),recording.id,filepath,filename_for_asr,start_time), kwargs={'language':language,'min_speakers':min_s,'max_speakers':max_s})
        else:
            start_time=datetime.utcnow(); thread=threading.Thread(target=transcribe_audio_task, args=(app.app_context(),recording.id,filepath,filename_for_asr,start_time))
        thread.start(); return jsonify({'success':True,'message':'Transcription reprocessing started','recording':recording.to_dict()})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error reprocessing transcription for recording {recording_id}: {e}")
        return jsonify({'error':str(e)}),500

@recordings_bp.route('/recording/<int:recording_id>/reprocess_summary', methods=['POST'])
@login_required
def reprocess_summary(recording_id):
    try:
        from flask import current_app as app
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to reprocess this recording'}),403
        if not recording.transcription or len(recording.transcription.strip())<10: return jsonify({'error':'No valid transcription available for summary generation'}),400
        if recording.status in ['PROCESSING','SUMMARIZING']: return jsonify({'error':'Recording is already being processed'}),400
        if call_llm_completion is None: return jsonify({'error':'Summary service is not available (OpenRouter client not configured)'}),503
        recording.summary=None; recording.status='SUMMARIZING'; Event.query.filter_by(recording_id=recording_id).delete(); db.session.commit(); db.session.refresh(recording)
        def _task(ctx, rid): generate_summary_only_task(ctx, rid)
        thread=threading.Thread(target=_task, args=(app.app_context(), recording.id)); thread.start();
        return jsonify({'success':True,'message':'Summary reprocessing started','recording':recording.to_dict()})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error reprocessing summary for recording {recording_id}: {e}")
        return jsonify({'error':str(e)}),500

@recordings_bp.route('/recording/<int:recording_id>/reset_status', methods=['POST'])
@login_required
def reset_status(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to modify this recording'}),403
        if recording.status in ['PROCESSING','SUMMARIZING','FAILED']:
            recording.status='FAILED'; recording.error_message='Manually reset from stuck or failed state.'; db.session.commit(); return jsonify({'success':True,'message':'Recording status has been reset.','recording':recording.to_dict()})
        return jsonify({'error':f'Recording is not in a state that can be reset. Current status: {recording.status}'}),400
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error resetting status for recording {recording_id}: {e}")
        return jsonify({'error':str(e)}),500


# --- Audio & Delete ---
@recordings_bp.route('/audio/<int:recording_id>')
@login_required
def get_audio(recording_id):
    try:
        recording=db.session.get(Recording, recording_id)
        if not recording or not recording.audio_path: return jsonify({'error':'Recording or audio file not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to access this audio file'}),403
        
        # Resolve audio path - convert relative to absolute using upload folder from config
        audio_path = recording.audio_path
        if not os.path.isabs(audio_path):
            # Relative path - resolve against upload folder
            audio_path = os.path.join(current_app.config['UPLOAD_FOLDER'], os.path.basename(audio_path))
        
        if not os.path.exists(audio_path): return jsonify({'error':'Audio file missing from server'}),404
        return send_file(audio_path)
    except Exception as e:
        current_app.logger.error(f"Error serving audio for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':'An unexpected error occurred.'}),500

@recordings_bp.route('/share/audio/<string:public_id>')
def get_shared_audio(public_id):
    try:
        share=Share.query.filter_by(public_id=public_id).first_or_404(); recording=share.recording
        if not recording or not recording.audio_path: return jsonify({'error':'Recording or audio file not found'}),404
        
        # Resolve audio path - convert relative to absolute using upload folder from config
        audio_path = recording.audio_path
        if not os.path.isabs(audio_path):
            # Relative path - resolve against upload folder
            audio_path = os.path.join(current_app.config['UPLOAD_FOLDER'], os.path.basename(audio_path))
        
        if not os.path.exists(audio_path): return jsonify({'error':'Audio file missing from server'}),404
        return send_file(audio_path)
    except Exception as e:
        current_app.logger.error(f"Error serving shared audio for public_id {public_id}: {e}", exc_info=True)
        return jsonify({'error':'An unexpected error occurred.'}),500

@recordings_bp.route('/recording/<int:recording_id>', methods=['DELETE'])
@login_required
def delete_recording(recording_id):
    try:
        from flask import current_app as app
        recording=db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error':'Recording not found'}),404
        if recording.user_id and recording.user_id!=current_user.id: return jsonify({'error':'You do not have permission to delete this recording'}),403
        try:
            if recording.audio_path and os.path.exists(recording.audio_path): os.remove(recording.audio_path)
        except Exception as e:
            current_app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")
        attachments=Attachment.query.filter_by(recording_id=recording_id).all()
        for attachment in attachments:
            try:
                if os.path.exists(attachment.file_path): os.remove(attachment.file_path)
            except Exception as e:
                current_app.logger.error(f"Error deleting attachment file {attachment.file_path}: {e}")
        if os.environ.get('ENABLE_INQUIRE_MODE') == 'true':
            chunk_count=TranscriptChunk.query.filter_by(recording_id=recording_id).count()
            if chunk_count>0: current_app.logger.info(f"Deleting {chunk_count} transcript chunks with embeddings for recording {recording_id}")
        db.session.delete(recording); db.session.commit()
        return jsonify({'success':True})
    except Exception as e:
        db.session.rollback(); current_app.logger.error(f"Error deleting recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error':'An unexpected error occurred while deleting.'}),500

